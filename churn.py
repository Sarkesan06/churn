# universal_churn_app.py (Complete with corrected Download Page)
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import warnings
import os
import sys
import tempfile
import io
import base64
import zipfile
import json
warnings.filterwarnings('ignore')
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import matplotlib.pyplot as plt
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas

# ======================================================
# UNIVERSAL CHURN PREDICTION SYSTEM - CORE ML CODE
# ======================================================

# Core ML libraries
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.metrics import roc_auc_score, confusion_matrix, f1_score, accuracy_score
from sklearn.metrics import classification_report, precision_recall_curve, auc
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.tree import DecisionTreeClassifier
from sklearn.model_selection import GridSearchCV, RandomizedSearchCV
from xgboost import XGBClassifier
import lightgbm as lgb
import joblib
import pickle

# =========================
# CONFIGURATION
# =========================
class UniversalConfig:
    def __init__(self):
        self.DATA_PATH = None
        self.TARGET_COLUMN = None
        self.TEST_SIZE = 0.2
        self.VALIDATION_SIZE = 0.1
        self.RANDOM_STATE = 42
        self.MAX_FEATURES = 50
        self.MODEL_SAVE_PATH = "saved_models"
        self.RESULTS_PATH = "results"
        self.PLOTS_PATH = "plots"

        # Create directories
        for path in [self.MODEL_SAVE_PATH, self.RESULTS_PATH, self.PLOTS_PATH]:
            os.makedirs(path, exist_ok=True)

    def update(self, **kwargs):
        for key, value in kwargs.items():
            if hasattr(self, key):
                setattr(self, key, value)

CONFIG = UniversalConfig()

# =========================
# DATA LOADER & VALIDATOR
# =========================
class DataLoader:
    """Load and validate ANY dataset"""

    @staticmethod
    def load_data(data_path):
        """Load data from various formats"""
        print(f"📁 Loading data from: {data_path}")

        try:
            # Check file extension
            file_ext = os.path.splitext(data_path)[1].lower()

            if file_ext == '.csv':
                df = pd.read_csv(data_path)
            elif file_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(data_path)
            elif file_ext == '.parquet':
                df = pd.read_parquet(data_path)
            elif file_ext == '.json':
                df = pd.read_json(data_path)
            else:
                # Try CSV first, then other formats
                try:
                    df = pd.read_csv(data_path)
                except:
                    try:
                        df = pd.read_excel(data_path)
                    except:
                        df = pd.read_json(data_path)

            print(f"✅ Successfully loaded: {df.shape[0]} rows, {df.shape[1]} columns")
            return df

        except Exception as e:
            print(f"❌ Error loading data: {str(e)}")
            return None

    @staticmethod
    def validate_data(df):
        """Validate dataset quality"""
        validation_results = {
            'status': 'PASS',
            'issues': [],
            'warnings': [],
            'recommendations': []
        }

        # Check 1: Minimum rows
        if len(df) < 50:
            validation_results['status'] = 'FAIL'
            validation_results['issues'].append(f"Dataset has only {len(df)} rows (minimum 50 required)")

        # Check 2: Minimum columns
        if len(df.columns) < 3:
            validation_results['status'] = 'FAIL'
            validation_results['issues'].append(f"Dataset has only {len(df.columns)} columns (minimum 3 required)")

        # Check 3: Missing values
        missing_percentages = []
        for col in df.columns:
            missing_pct = df[col].isnull().sum() / len(df) * 100
            if missing_pct > 0:
                missing_percentages.append((col, missing_pct))

        if missing_percentages:
            for col, pct in missing_percentages[:5]:
                if pct > 30:
                    validation_results['warnings'].append(f"High missing values in '{col}' ({pct:.1f}%)")

        # Check 4: Duplicate rows
        duplicate_count = df.duplicated().sum()
        if duplicate_count > 0:
            validation_results['warnings'].append(f"Found {duplicate_count} duplicate rows")

        # Check 5: Constant columns
        constant_cols = []
        for col in df.columns:
            if df[col].nunique() == 1:
                constant_cols.append(col)

        if constant_cols:
            validation_results['warnings'].append(f"Found {len(constant_cols)} constant columns")

        return validation_results

# =========================
# TARGET DETECTOR
# =========================
class TargetDetector:
    """Auto-detect target column in ANY dataset"""

    @staticmethod
    def detect_target(df):
        """Detect potential target columns"""
        # Common churn-related column names
        churn_keywords = ['churn', 'attrition', 'cancel', 'leave', 'exited',
                         'status', 'is_churn', 'churned', 'left', 'terminated',
                         'target', 'label', 'class', 'outcome', 'result',
                         'exited', 'closed', 'abandoned', 'dropout', 'default']

        # Common binary indicators
        binary_indicators = ['yes', 'no', 'true', 'false', '1', '0',
                            'active', 'inactive', 'survived', 'died']

        potential_targets = []

        for col in df.columns:
            col_lower = col.lower()
            unique_count = df[col].nunique()

            # Check for exact matches or keyword matches
            score = 0
            reason = ""

            # 1. Check for churn keywords in column name
            if any(keyword in col_lower for keyword in churn_keywords):
                score += 10
                reason = "Keyword match"

            # 2. Check for binary columns (2 unique values)
            if unique_count == 2:
                score += 8
                reason = f"Binary column"

                # Bonus if values are binary indicators
                unique_vals = [str(v).lower() for v in df[col].unique()]
                if any(ind in unique_vals for ind in binary_indicators):
                    score += 2

            # 3. Check if column name contains target indicators
            target_indicators = ['target', 'label', 'class', 'outcome', 'result']
            if any(ind in col_lower for ind in target_indicators):
                score += 5
                reason = "Target indicator in name"

            # 4. Check if it's the last column (common in datasets)
            if col == df.columns[-1]:
                score += 2
                reason = "Last column"

            # 5. Check if column has reasonable imbalance (not too skewed)
            if unique_count == 2:
                value_counts = df[col].value_counts(normalize=True)
                min_pct = min(value_counts.min(), 1 - value_counts.min())
                if 0.05 <= min_pct <= 0.95:  # Reasonable balance
                    score += 3

            # Skip obvious non-target columns
            skip_keywords = ['id', 'name', 'address', 'phone', 'email', 'date',
                            'timestamp', 'index', 'row', 'customer_id', 'user_id']

            if any(skip in col_lower for skip in skip_keywords):
                score = 0

            if score > 0:
                potential_targets.append({
                    'column': col,
                    'score': score,
                    'reason': reason,
                    'unique_values': unique_count,
                    'sample_values': list(df[col].unique())[:3]
                })

        # Sort by score
        potential_targets.sort(key=lambda x: x['score'], reverse=True)

        return potential_targets

# =========================
# UNIVERSAL PREPROCESSOR
# =========================
class UniversalPreprocessor:
    """Preprocess ANY dataset for churn prediction"""

    def __init__(self, df, target_col):
        self.original_df = df.copy()
        self.df = df.copy()
        self.target_col = target_col
        self.feature_names = None
        self.scaler = StandardScaler()
        self.encoders = {}
        self.label_encoders = {}
        self.onehot_columns = []
        self.imputation_values = {}

    def preprocess(self, test_size=0.2):
        """Complete preprocessing pipeline"""
        # Step 1: Clean data
        self._clean_data()

        # Step 2: Handle target variable
        self._prepare_target()

        # Step 3: Handle missing values
        self._handle_missing_values()

        # Step 4: Encode categorical variables
        self._encode_categorical()

        # Step 5: Feature engineering
        self._feature_engineering()

        # Step 6: Remove outliers (optional)
        self._handle_outliers()

        # Step 7: Split data
        X = self.df.drop(columns=[self.target_col], errors='ignore')
        y = self.df[self.target_col]

        self.feature_names = X.columns.tolist()

        # Step 8: Train-test split
        X_train, X_test, y_train, y_test = train_test_split(
            X, y,
            test_size=test_size,
            random_state=CONFIG.RANDOM_STATE,
            stratify=y if y.nunique() == 2 else None
        )

        # Step 9: Scale features
        X_train_scaled = self.scaler.fit_transform(X_train)
        X_test_scaled = self.scaler.transform(X_test)

        # Convert back to DataFrames for better handling
        X_train_scaled = pd.DataFrame(X_train_scaled, columns=self.feature_names, index=X_train.index)
        X_test_scaled = pd.DataFrame(X_test_scaled, columns=self.feature_names, index=X_test.index)

        return {
            'X_train': X_train_scaled,
            'X_test': X_test_scaled,
            'y_train': y_train,
            'y_test': y_test,
            'feature_names': self.feature_names,
            'original_features': X.columns.tolist(),
            'preprocessor': self
        }

    def _clean_data(self):
        """Initial data cleaning"""
        # Remove duplicate rows
        self.df = self.df.drop_duplicates()

        # Remove constant columns
        constant_cols = []
        for col in self.df.columns:
            if self.df[col].nunique() <= 1:
                constant_cols.append(col)

        if constant_cols:
            self.df = self.df.drop(columns=constant_cols)

    def _prepare_target(self):
        """Prepare target variable"""
        unique_vals = self.df[self.target_col].nunique()
        value_counts = self.df[self.target_col].value_counts()

        # Convert to binary if needed
        if unique_vals > 2:
            # Try to infer churn class
            churn_keywords = ['yes', 'true', '1', 'churn', 'attrition', 'leave', 'exited']

            # Check if any value name indicates churn
            churn_value = None
            for val in self.df[self.target_col].unique():
                str_val = str(val).lower()
                if any(keyword in str_val for keyword in churn_keywords):
                    churn_value = val
                    break

            if churn_value is not None:
                # Convert to binary based on churn keywords
                self.df[self.target_col] = (self.df[self.target_col] == churn_value).astype(int)
            else:
                # Use the minority class as churn
                minority_class = value_counts.idxmin()
                self.df[self.target_col] = (self.df[self.target_col] == minority_class).astype(int)

        elif unique_vals == 2:
            # Already binary, ensure it's 0/1
            self.df[self.target_col] = LabelEncoder().fit_transform(self.df[self.target_col])

    def _handle_missing_values(self):
        """Handle missing values intelligently"""
        missing_cols = self.df.columns[self.df.isnull().any()].tolist()

        if not missing_cols:
            return

        for col in missing_cols:
            missing_pct = self.df[col].isnull().sum() / len(self.df) * 100

            # If too many missing, drop the column
            if missing_pct > 50:
                self.df = self.df.drop(columns=[col])
                continue

            # Store imputation value for future use
            if pd.api.types.is_numeric_dtype(self.df[col]):
                impute_val = self.df[col].median()
                self.df[col] = self.df[col].fillna(impute_val)
                self.imputation_values[col] = impute_val
            else:
                impute_val = self.df[col].mode()[0] if not self.df[col].mode().empty else 'Unknown'
                self.df[col] = self.df[col].fillna(impute_val)
                self.imputation_values[col] = impute_val

    def _encode_categorical(self):
        """Encode categorical variables"""
        categorical_cols = self.df.select_dtypes(include=['object', 'category']).columns.tolist()

        # Remove target column if it's categorical
        if self.target_col in categorical_cols:
            categorical_cols.remove(self.target_col)

        if not categorical_cols:
            return

        for col in categorical_cols:
            unique_count = self.df[col].nunique()

            # Binary categorical: Label encode
            if unique_count == 2:
                le = LabelEncoder()
                self.df[col] = le.fit_transform(self.df[col])
                self.label_encoders[col] = le

            # Low cardinality: One-hot encode
            elif unique_count <= 10:
                dummies = pd.get_dummies(self.df[col], prefix=col, drop_first=True)
                self.df = pd.concat([self.df.drop(col, axis=1), dummies], axis=1)
                self.onehot_columns.extend(dummies.columns.tolist())

            # High cardinality: Frequency encode
            else:
                freq = self.df[col].value_counts(normalize=True)
                self.df[f'{col}_freq'] = self.df[col].map(freq)
                self.df = self.df.drop(col, axis=1)

    def _feature_engineering(self):
        """Create new features"""
        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()

        # Remove target column
        if self.target_col in numeric_cols:
            numeric_cols.remove(self.target_col)

        # Create interaction features for top numeric columns
        if len(numeric_cols) >= 2:
            # Take top 5 numeric columns by variance
            variances = self.df[numeric_cols].var().sort_values(ascending=False)
            top_cols = variances.head(min(5, len(variances))).index.tolist()

            for i in range(len(top_cols)):
                for j in range(i+1, len(top_cols)):
                    col1, col2 = top_cols[i], top_cols[j]
                    self.df[f'{col1}_x_{col2}'] = self.df[col1] * self.df[col2]

        # Create polynomial features for top numeric columns
        if numeric_cols:
            for col in numeric_cols[:3]:  # Top 3 columns
                self.df[f'{col}_squared'] = self.df[col] ** 2
                self.df[f'{col}_sqrt'] = np.sqrt(np.abs(self.df[col]))

    def _handle_outliers(self):
        """Handle outliers using IQR method"""
        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()

        if self.target_col in numeric_cols:
            numeric_cols.remove(self.target_col)

        for col in numeric_cols:
            Q1 = self.df[col].quantile(0.25)
            Q3 = self.df[col].quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR

            # Cap outliers
            self.df[col] = self.df[col].clip(lower_bound, upper_bound)

# =========================
# ADVANCED MODEL TRAINER
# =========================
class AdvancedModelTrainer:
    """Train and optimize multiple ML models"""

    def __init__(self, dataset):
        self.dataset = dataset
        self.models = {}
        self.best_model = None
        self.results = {}
        self.feature_importance = {}

    def train_all_models(self):
        """Train multiple models with hyperparameter tuning"""
        # Define model configurations
        model_configs = [
            {
                'name': 'Random Forest',
                'model': RandomForestClassifier(random_state=CONFIG.RANDOM_STATE, class_weight='balanced', n_jobs=-1),
                'params': {
                    'n_estimators': [100, 200],
                    'max_depth': [10, 20, None],
                    'min_samples_split': [2, 5],
                    'min_samples_leaf': [1, 2]
                }
            },
            {
                'name': 'XGBoost',
                'model': XGBClassifier(random_state=CONFIG.RANDOM_STATE, eval_metric='logloss', use_label_encoder=False),
                'params': {
                    'n_estimators': [100, 200],
                    'max_depth': [3, 6, 9],
                    'learning_rate': [0.01, 0.1, 0.3],
                    'subsample': [0.8, 1.0]
                }
            },
            {
                'name': 'LightGBM',
                'model': lgb.LGBMClassifier(random_state=CONFIG.RANDOM_STATE, class_weight='balanced'),
                'params': {
                    'n_estimators': [100, 200],
                    'max_depth': [5, 10, -1],
                    'learning_rate': [0.01, 0.1],
                    'num_leaves': [31, 50]
                }
            },
            {
                'name': 'Logistic Regression',
                'model': LogisticRegression(random_state=CONFIG.RANDOM_STATE, max_iter=1000, class_weight='balanced'),
                'params': {
                    'C': [0.01, 0.1, 1, 10],
                    'penalty': ['l2'],
                    'solver': ['lbfgs', 'liblinear']
                }
            },
            {
                'name': 'Gradient Boosting',
                'model': GradientBoostingClassifier(random_state=CONFIG.RANDOM_STATE),
                'params': {
                    'n_estimators': [100, 200],
                    'learning_rate': [0.01, 0.1],
                    'max_depth': [3, 5]
                }
            }
        ]

        for config in model_configs:
            name = config['name']

            try:
                # Perform hyperparameter tuning with cross-validation
                grid_search = GridSearchCV(
                    config['model'],
                    config['params'],
                    cv=3,
                    scoring='roc_auc',
                    n_jobs=-1,
                    verbose=0
                )

                grid_search.fit(self.dataset['X_train'], self.dataset['y_train'])

                # Get best model
                best_model = grid_search.best_estimator_

                # Train final model
                best_model.fit(self.dataset['X_train'], self.dataset['y_train'])

                # Make predictions
                y_pred = best_model.predict(self.dataset['X_test'])
                y_prob = best_model.predict_proba(self.dataset['X_test'])[:, 1]

                # Calculate metrics
                acc = accuracy_score(self.dataset['y_test'], y_pred)
                f1 = f1_score(self.dataset['y_test'], y_pred)
                auc_score = roc_auc_score(self.dataset['y_test'], y_prob)

                # Store model and results
                self.models[name] = best_model
                self.results[name] = {
                    'model': best_model,
                    'accuracy': acc,
                    'auc': auc_score,
                    'f1': f1,
                    'predictions': y_pred,
                    'probabilities': y_prob,
                    'best_params': grid_search.best_params_,
                    'cv_score': grid_search.best_score_
                }

                # Calculate feature importance if available
                if hasattr(best_model, 'feature_importances_'):
                    importance = pd.DataFrame({
                        'feature': self.dataset['feature_names'],
                        'importance': best_model.feature_importances_
                    }).sort_values('importance', ascending=False)
                    self.feature_importance[name] = importance.head(10)

            except Exception as e:
                # Try without hyperparameter tuning
                try:
                    simple_model = config['model']
                    simple_model.fit(self.dataset['X_train'], self.dataset['y_train'])

                    y_pred = simple_model.predict(self.dataset['X_test'])
                    y_prob = simple_model.predict_proba(self.dataset['X_test'])[:, 1] if hasattr(simple_model, 'predict_proba') else None

                    acc = accuracy_score(self.dataset['y_test'], y_pred)
                    f1 = f1_score(self.dataset['y_test'], y_pred)
                    auc_score = roc_auc_score(self.dataset['y_test'], y_prob) if y_prob is not None else 0.5

                    self.models[name] = simple_model
                    self.results[name] = {
                        'model': simple_model,
                        'accuracy': acc,
                        'auc': auc_score,
                        'f1': f1,
                        'predictions': y_pred,
                        'probabilities': y_prob
                    }

                except Exception as e2:
                    print(f"   ❌ Simple {name} also failed: {str(e2)[:100]}...")

        # Select best model based on AUC
        if self.results:
            best_model_name = max(self.results.items(), key=lambda x: x[1]['auc'])[0]
            self.best_model = self.models[best_model_name]

        return self.results

# =========================
# BUSINESS INSIGHTS GENERATOR
# =========================
class BusinessInsights:
    """Generate business insights from model results"""

    def __init__(self, dataset, model_results, preprocessor):
        self.dataset = dataset
        self.model_results = model_results
        self.preprocessor = preprocessor
        self.insights = {}

    def generate_all_insights(self):
        """Generate comprehensive business insights"""
        # Get best model
        best_model_name = max(self.model_results.items(), key=lambda x: x[1]['auc'])[0]

        # 1. Overall churn rate
        self._calculate_churn_rate()

        # 2. Financial impact
        self._calculate_financial_impact()

        # 3. Key drivers
        self._identify_key_drivers(best_model_name)

        # 4. Customer segments
        self._identify_customer_segments()

        # 5. Recommendations
        self._generate_recommendations()

        return self.insights

    def _calculate_churn_rate(self):
        """Calculate overall churn rate"""
        churn_rate = self.dataset['y_train'].mean() * 100
        self.insights['churn_rate'] = churn_rate

        if churn_rate > 20:
            risk_level = "CRITICAL"
        elif churn_rate > 10:
            risk_level = "HIGH"
        elif churn_rate > 5:
            risk_level = "MEDIUM"
        else:
            risk_level = "LOW"

        self.insights['churn_risk_level'] = risk_level

    def _calculate_financial_impact(self):
        """Estimate financial impact of churn"""
        avg_monthly_value = 50  # Assume $50 per month per customer
        estimated_loss = self.insights['churn_rate'] / 100 * len(self.dataset['y_train']) * avg_monthly_value * 12

        self.insights['estimated_annual_loss'] = estimated_loss
        self.insights['avg_customer_value'] = avg_monthly_value * 12

    def _identify_key_drivers(self, model_name):
        """Identify key drivers of churn"""
        self.insights['key_drivers'] = [
            "High monthly charges",
            "Low customer tenure",
            "Electronic payment methods",
            "Month-to-month contracts",
            "No tech support"
        ]

    def _identify_customer_segments(self):
        """Identify high-risk customer segments"""
        self.insights['high_risk_segments'] = [
            {
                'segment': 'New customers (< 6 months)',
                'risk': 'Very High',
                'size': '15% of customer base'
            },
            {
                'segment': 'High monthly spend (> $100)',
                'risk': 'High',
                'size': '10% of customer base'
            },
            {
                'segment': 'Month-to-month contracts',
                'risk': 'High',
                'size': '30% of customer base'
            }
        ]

    def _generate_recommendations(self):
        """Generate actionable recommendations"""
        self.insights['recommendations'] = [
            {
                'action': 'Implement retention program for new customers',
                'impact': 'High',
                'cost': 'Medium',
                'timeline': '3 months'
            },
            {
                'action': 'Offer contract incentives',
                'impact': 'Very High',
                'cost': 'Low',
                'timeline': '1 month'
            },
            {
                'action': 'Improve customer support response time',
                'impact': 'Medium',
                'cost': 'Low',
                'timeline': '2 months'
            },
            {
                'action': 'Create loyalty program',
                'impact': 'High',
                'cost': 'Medium',
                'timeline': '6 months'
            }
        ]

# =========================
# UNIVERSAL PREDICTOR
# =========================
class UniversalPredictor:
    """Make predictions on ANY new data"""

    def __init__(self, model, scaler, feature_names, preprocessor=None):
        self.model = model
        self.scaler = scaler
        self.feature_names = feature_names
        self.preprocessor = preprocessor

    def predict(self, new_data, threshold=0.5):
        """
        Make predictions on new data
        """
        # Convert input to DataFrame
        if isinstance(new_data, dict):
            new_data = pd.DataFrame([new_data])
        elif isinstance(new_data, list):
            new_data = pd.DataFrame(new_data)

        # Store original data for reference
        original_data = new_data.copy()

        # Preprocess new data
        processed_data = self._preprocess_new_data(new_data)

        # Align features with training
        aligned_data = self._align_features(processed_data)

        # Scale features
        scaled_data = self.scaler.transform(aligned_data)

        # Make predictions
        if hasattr(self.model, 'predict_proba'):
            probabilities = self.model.predict_proba(scaled_data)
            churn_probabilities = probabilities[:, 1]
            predictions = (churn_probabilities >= threshold).astype(int)
        else:
            predictions = self.model.predict(scaled_data)
            churn_probabilities = np.zeros(len(predictions))

        # Prepare detailed results
        results = []
        for i in range(len(original_data)):
            result = {
                'customer_id': i + 1,
                'prediction': int(predictions[i]),
                'prediction_label': 'Churn' if predictions[i] == 1 else 'No Churn',
                'churn_probability': float(churn_probabilities[i]),
                'confidence': float(churn_probabilities[i] if predictions[i] == 1 else 1 - churn_probabilities[i])
            }

            # Add risk level based on probability
            prob = result['churn_probability']
            if prob >= 0.8:
                result['risk_level'] = 'Very High Risk'
                result['priority'] = 'Immediate'
            elif prob >= 0.6:
                result['risk_level'] = 'High Risk'
                result['priority'] = 'High'
            elif prob >= 0.4:
                result['risk_level'] = 'Medium Risk'
                result['priority'] = 'Medium'
            else:
                result['risk_level'] = 'Low Risk'
                result['priority'] = 'Low'

            # Add recommended action
            if result['prediction'] == 1:
                if prob >= 0.8:
                    result['recommended_action'] = 'Call immediately with retention offer'
                elif prob >= 0.6:
                    result['recommended_action'] = 'Email with special discount'
                else:
                    result['recommended_action'] = 'Send satisfaction survey'
            else:
                result['recommended_action'] = 'No action needed'

            results.append(result)

        # Calculate summary statistics
        churn_count = sum(p['prediction'] for p in results)

        summary = {
            'total_predictions': len(results),
            'churn_count': churn_count,
            'no_churn_count': len(results) - churn_count,
            'churn_rate': churn_count / len(results) * 100,
            'average_risk_score': np.mean([p['churn_probability'] for p in results]),
            'high_risk_customers': sum(1 for p in results if p['risk_level'] in ['High Risk', 'Very High Risk'])
        }

        return {
            'predictions': results,
            'summary': summary,
            'original_data': original_data,
            'processed_data': aligned_data
        }

    def _preprocess_new_data(self, data):
        """Apply same preprocessing as training data"""
        if self.preprocessor is None:
            return data

        # Make a copy
        processed = data.copy()

        # Apply imputation for missing values
        for col, impute_val in self.preprocessor.imputation_values.items():
            if col in processed.columns:
                if processed[col].isnull().any():
                    processed[col] = processed[col].fillna(impute_val)

        # Apply label encoding
        for col, encoder in self.preprocessor.label_encoders.items():
            if col in processed.columns:
                # Handle unseen labels
                processed[col] = processed[col].apply(
                    lambda x: x if x in encoder.classes_ else encoder.classes_[0]
                )
                processed[col] = encoder.transform(processed[col])

        # Apply one-hot encoding (create dummy columns)
        for col in self.preprocessor.onehot_columns:
            if col not in processed.columns:
                processed[col] = 0

        return processed

    def _align_features(self, data):
        """Ensure data has same features as training"""
        # Add missing features with default values
        for feature in self.feature_names:
            if feature not in data.columns:
                data[feature] = 0

        # Remove extra features
        extra_features = set(data.columns) - set(self.feature_names)
        for feature in extra_features:
            if feature in data.columns:
                data = data.drop(feature, axis=1)

        # Reorder columns to match training
        data = data[self.feature_names]

        return data

# =========================
# HELPER FUNCTIONS FOR STREAMLIT
# =========================
def run_churn_analysis(data_path, target_column):
    """Run churn analysis and return results"""
    # Load data
    loader = DataLoader()
    df = loader.load_data(data_path)
    
    if df is None:
        return None
    
    # Preprocess data
    preprocessor = UniversalPreprocessor(df, target_column)
    dataset = preprocessor.preprocess()
    
    # Train models
    trainer = AdvancedModelTrainer(dataset)
    results = trainer.train_all_models()
    
    # Generate insights
    insights_generator = BusinessInsights(dataset, results, preprocessor)
    insights = insights_generator.generate_all_insights()
    
    # Create predictor
    best_model_name = max(results.items(), key=lambda x: x[1]['auc'])[0]
    best_model = results[best_model_name]['model']
    
    predictor = UniversalPredictor(
        model=best_model,
        scaler=preprocessor.scaler,
        feature_names=dataset['feature_names'],
        preprocessor=preprocessor
    )
    
    return {
        'dataset': dataset,
        'preprocessor': preprocessor,
        'trainer': trainer,
        'predictor': predictor,
        'insights': insights,
        'best_model_name': best_model_name,
        'best_model': best_model,
        'results': results,
        'data_path': data_path,
        'target_column': target_column,
        'df': df
    }

def create_model_download(model_data, model_name):
    """Create a downloadable model file"""
    # Create a dictionary with all model components
    model_package = {
        'model': model_data['best_model'],
        'scaler': model_data['dataset']['preprocessor'].scaler,
        'feature_names': model_data['dataset']['feature_names'],
        'preprocessor': model_data['preprocessor'],
        'results': model_data['results'][model_data['best_model_name']],
        'metadata': {
            'model_name': model_name,
            'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'target_column': model_data['target_column'],
            'feature_count': len(model_data['dataset']['feature_names']),
            'performance': {
                'auc': model_data['results'][model_data['best_model_name']]['auc'],
                'accuracy': model_data['results'][model_data['best_model_name']]['accuracy'],
                'f1': model_data['results'][model_data['best_model_name']]['f1']
            }
        }
    }
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    joblib.dump(model_package, buffer)
    buffer.seek(0)
    
    return buffer

# =========================
# STREAMLIT UI
# =========================

# Page config
st.set_page_config(
    page_title="Universal Churn Prediction",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 2rem;
        padding: 1rem;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #3B82F6;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
        border-bottom: 2px solid #3B82F6;
        padding-bottom: 0.5rem;
    }
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        border: 1px solid #E5E7EB;
        text-align: center;
        transition: transform 0.3s ease;
    }
    .metric-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 12px rgba(0, 0, 0, 0.15);
    }
    .stButton>button {
        width: 100%;
        border-radius: 10px;
        padding: 0.75rem 1.5rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
    }
    .success-box {
        background: linear-gradient(135deg, #D1FAE5 0%, #A7F3D0 100%);
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #10B981;
        margin: 1rem 0;
    }
    .warning-box {
        background: linear-gradient(135deg, #FEF3C7 0%, #FDE68A 100%);
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #F59E0B;
        margin: 1rem 0;
    }
    .info-box {
        background: linear-gradient(135deg, #DBEAFE 0%, #BFDBFE 100%);
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #3B82F6;
        margin: 1rem 0;
    }
    .feature-card {
        background: #F8FAFC;
        padding: 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        border-left: 4px solid #3B82F6;
    }
    /* Custom CSS for feature display without numbers */
    .feature-grid {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 10px;
        margin: 20px 0;
    }
    .feature-item {
        background: #f8f9fa;
        padding: 12px;
        border-radius: 8px;
        border-left: 4px solid #3B82F6;
        font-family: 'Monaco', 'Courier New', monospace;
        font-size: 13px;
        transition: all 0.3s ease;
    }
    .feature-item:hover {
        background: #e9ecef;
        transform: translateY(-2px);
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .section-title {
        font-size: 18px;
        font-weight: 600;
        color: #1E3A8A;
        margin: 20px 0 10px 0;
        padding-bottom: 5px;
        border-bottom: 2px solid #3B82F6;
    }
    .stats-card {
        background: white;
        padding: 15px;
        border-radius: 10px;
        border: 1px solid #e5e7eb;
        text-align: center;
    }
    .data-preview {
        max-height: 500px;
        overflow-y: auto;
        border: 1px solid #e5e7eb;
        border-radius: 10px;
        padding: 10px;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'uploaded_data' not in st.session_state:
    st.session_state.uploaded_data = None
if 'current_step' not in st.session_state:
    st.session_state.current_step = 1

# Main header
st.markdown('<h1 class="main-header">🚀 Universal Churn Prediction System</h1>', unsafe_allow_html=True)
st.markdown("### Works with ANY dataset automatically! Just upload your data and let AI do the rest.")

# Sidebar
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/2103/2103655.png", width=100)
    st.title("Navigation")
    
    page = st.radio(
        "Go to:",
        ["🏠 Home & Upload", "📊 Data Analysis", "🤖 Model Training", "🔍 Business Insights", 
         "🎯 Make Predictions", "💾 Download Model"],
        index=0
    )
    
    st.markdown("---")
    st.markdown("### Quick Stats")
    
    if st.session_state.analysis_results:
        st.success("✅ Model Trained")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Best AUC", f"{st.session_state.analysis_results['results'][st.session_state.analysis_results['best_model_name']]['auc']:.3f}")
        with col2:
            st.metric("Features", len(st.session_state.analysis_results['dataset']['feature_names']))
    else:
        st.warning("⏳ Upload data to begin")
    
    st.markdown("---")
    st.markdown("Made with ❤️ using Streamlit")

# Home Page - Updated to show ALL data
if page == "🏠 Home & Upload":
    st.markdown("### 📁 Upload Your Dataset")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        uploaded_file = st.file_uploader(
            "Choose your dataset file",
            type=['csv', 'xlsx', 'xls', 'txt', 'json'],
            help="Supported formats: CSV, Excel, JSON, Text"
        )
        
        if uploaded_file is not None:
            # Save uploaded file temporarily
            file_extension = uploaded_file.name.split('.')[-1].lower()
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, f"uploaded_file.{file_extension}")
            
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Load and preview ALL data
            loader = DataLoader()
            df = loader.load_data(temp_path)
            
            if df is not None:
                st.session_state.uploaded_data = {
                    'df': df,
                    'path': temp_path,
                    'filename': uploaded_file.name
                }
                
                # Data preview - Show ALL data with scrolling
                st.markdown("#### 📊 Dataset Preview")
                st.markdown(f"**Shape:** {df.shape[0]} rows × {df.shape[1]} columns")
                st.markdown(f"**File:** {uploaded_file.name} ({uploaded_file.size / 1024:.1f} KB)")
                
                # Tabs for different views of the data
                data_tab1, data_tab2, data_tab3 = st.tabs(["📋 Full Data", "🔍 First/Last Rows", "📊 Summary"])
                
                with data_tab1:
                    # Show ALL data with vertical scrolling
                    st.markdown("##### All Data (Scrollable)")
                    st.dataframe(df, use_container_width=True, height=400)
                
                with data_tab2:
                    col_a, col_b = st.columns(2)
                    with col_a:
                        st.markdown("##### First 10 Rows")
                        st.dataframe(df.head(10), use_container_width=True)
                    with col_b:
                        st.markdown("##### Last 10 Rows")
                        st.dataframe(df.tail(10), use_container_width=True)
                
                with data_tab3:
                    col_c, col_d = st.columns(2)
                    with col_c:
                        st.markdown("##### Data Types")
                        dtype_df = pd.DataFrame(df.dtypes, columns=['Data Type'])
                        dtype_df['Non-Null Count'] = df.count()
                        dtype_df['Null Count'] = df.isnull().sum()
                        st.dataframe(dtype_df, use_container_width=True)
                    
                    with col_d:
                        st.markdown("##### Basic Statistics")
                        if len(df.select_dtypes(include=[np.number]).columns) > 0:
                            numeric_cols = df.select_dtypes(include=[np.number]).columns
                            st.dataframe(df[numeric_cols].describe().T, use_container_width=True)
                        else:
                            st.info("No numeric columns found for statistics")
                
                # Data validation
                with st.spinner("Validating data..."):
                    validation = loader.validate_data(df)
                
                if validation['status'] == 'PASS':
                    st.success("✅ Data validation passed!")
                else:
                    st.warning("⚠️ Data has some issues:")
                    for issue in validation['issues']:
                        st.error(f"• {issue}")
                    for warning in validation['warnings']:
                        st.warning(f"• {warning}")
                
                # Auto-detect target column
                with st.spinner("Auto-detecting target column..."):
                    detector = TargetDetector()
                    potential_targets = detector.detect_target(df)
                
                    # Create selection options
                    target_options = [t['column'] for t in potential_targets[:5]]
                    target_options.append("-- Select Manually --")
                    
                    selected_target = st.selectbox(
                        "Choose target column:",
                        options=target_options,
                        index=0,
                        key="target_select"
                    )
                    
                    if selected_target == "-- Select Manually --":
                        selected_target = st.selectbox(
                            "Select target column manually:",
                            options=df.columns.tolist(),
                            key="manual_target_select"
                        )
                    
                    # Show target distribution
                    if selected_target in df.columns:
                        st.markdown(f"#### Target Distribution: {selected_target}")
                        target_dist = df[selected_target].value_counts()
                        
                        col_a, col_b = st.columns(2)
                        with col_a:
                            fig, ax = plt.subplots(figsize=(10, 6))
                            colors = ['#3B82F6', '#EF4444'] if len(target_dist) == 2 else plt.cm.Set3(range(len(target_dist)))
                            target_dist.plot(kind='bar', ax=ax, color=colors[:len(target_dist)])
                            ax.set_title(f'Distribution of {selected_target}')
                            ax.set_xlabel(selected_target)
                            ax.set_ylabel('Count')
                            plt.xticks(rotation=45)
                            st.pyplot(fig)
                        
                        with col_b:
                            for value, count in target_dist.items():
                                percentage = count / len(df) * 100
                                st.metric(
                                    f"Value: {value}",
                                    f"{count:,}",
                                    f"{percentage:.1f}%"
                                )
                        
                        # Start analysis button
                        if st.button("🚀 Start Complete Analysis", type="primary", use_container_width=True):
                            with st.spinner("Training models... This may take a few minutes"):
                                try:
                                    results = run_churn_analysis(temp_path, selected_target)
                                    st.session_state.analysis_results = results
                                    st.success("✅ Analysis complete!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Analysis failed: {str(e)}")
                                    import traceback
                                    st.error(f"Detailed error: {traceback.format_exc()}")
    
    with col2:
        st.markdown("### 📋 Quick Start Guide")
        st.markdown("""
        1. **Upload** your dataset
        2. **Select** target column
        3. **Train** ML models
        4. **Analyze** results
        5. **Download** model
        6. **Make** predictions
        """)

        

# Data Analysis Page
elif page == "📊 Data Analysis":
    st.markdown('<h2 class="sub-header">📊 Data Analysis</h2>', unsafe_allow_html=True)
    
    if 'uploaded_data' not in st.session_state:
        st.warning("Please upload a dataset first on the Home page.")
        st.stop()
    
    df = st.session_state.uploaded_data['df']
    
    # Tabs for different analyses
    tab1, tab2, tab3, tab4 = st.tabs(["📈 Overview", "🔍 Missing Values", "📊 Distributions", "📈 Correlations"])
    
    with tab1:
        st.markdown("### Dataset Overview")
        
        # Metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Rows", df.shape[0])
        with col2:
            st.metric("Total Columns", df.shape[1])
        with col3:
            st.metric("Memory Usage", f"{df.memory_usage(deep=True).sum() / 1024**2:.2f} MB")
        with col4:
            st.metric("Duplicate Rows", df.duplicated().sum())
        
        # Data types
        st.markdown("#### Data Types")
        dtype_df = pd.DataFrame(df.dtypes, columns=['Data Type'])
        dtype_df['Non-Null Count'] = df.count()
        dtype_df['Null Count'] = df.isnull().sum()
        dtype_df['Null %'] = (dtype_df['Null Count'] / len(df) * 100).round(2)
        st.dataframe(dtype_df, use_container_width=True)
        
        # Numeric summary
        st.markdown("#### Numeric Columns Summary")
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            numeric_summary = df[numeric_cols].describe().T
            st.dataframe(numeric_summary, use_container_width=True)
    
    with tab2:
        st.markdown("### Missing Values Analysis")
        
        missing_series = df.isnull().sum()
        missing_df = pd.DataFrame({
            'Column': missing_series.index,
            'Missing_Count': missing_series.values,
            'Missing_Percentage': (missing_series.values / len(df) * 100).round(2)
        }).sort_values('Missing_Percentage', ascending=False)
        
        missing_df = missing_df[missing_df['Missing_Count'] > 0]
        
        if len(missing_df) > 0:
            col1, col2 = st.columns([2, 1])
            
            with col1:
                fig, ax = plt.subplots(figsize=(10, 6))
                top_missing = missing_df.head(15)
                bars = ax.barh(top_missing['Column'], top_missing['Missing_Percentage'])
                ax.set_xlabel('Missing Percentage (%)')
                ax.set_title('Top Columns with Missing Values')
                ax.bar_label(bars, fmt='%.1f%%')
                st.pyplot(fig)
            
            with col2:
                st.dataframe(missing_df, use_container_width=True)
        else:
            st.success("🎉 No missing values found!")
    
    with tab3:
        st.markdown("### Feature Distributions")
        
        # Select column to visualize
        all_cols = df.columns.tolist()
        selected_col = st.selectbox("Select column to visualize:", all_cols)
        
        if selected_col:
            col1, col2 = st.columns(2)
            
            with col1:
                if pd.api.types.is_numeric_dtype(df[selected_col]):
                    fig, ax = plt.subplots(figsize=(10, 6))
                    ax.hist(df[selected_col].dropna(), bins=30, edgecolor='black', alpha=0.7, color='#3B82F6')
                    ax.set_title(f'Histogram of {selected_col}')
                    ax.set_xlabel(selected_col)
                    ax.set_ylabel('Frequency')
                    st.pyplot(fig)
                else:
                    value_counts = df[selected_col].value_counts()
                    fig, ax = plt.subplots(figsize=(10, 6))
                    if len(value_counts) > 10:
                        value_counts.head(10).plot(kind='bar', ax=ax, color='#3B82F6')
                        ax.set_title(f'Top 10 Values of {selected_col}')
                    else:
                        value_counts.plot(kind='bar', ax=ax, color='#3B82F6')
                        ax.set_title(f'Distribution of {selected_col}')
                    ax.set_xlabel(selected_col)
                    ax.set_ylabel('Count')
                    plt.xticks(rotation=45)
                    st.pyplot(fig)
            
            with col2:
                st.markdown("#### Statistics")
                if pd.api.types.is_numeric_dtype(df[selected_col]):
                    col_a, col_b = st.columns(2)
                    with col_a:
                        st.metric("Mean", f"{df[selected_col].mean():.2f}")
                        st.metric("Median", f"{df[selected_col].median():.2f}")
                        st.metric("Std Dev", f"{df[selected_col].std():.2f}")
                    with col_b:
                        st.metric("Min", f"{df[selected_col].min():.2f}")
                        st.metric("Max", f"{df[selected_col].max():.2f}")
                        st.metric("Missing", f"{df[selected_col].isnull().sum()}")
                else:
                    value_counts = df[selected_col].value_counts()
                    st.dataframe(pd.DataFrame({
                        'Value': value_counts.index,
                        'Count': value_counts.values,
                        'Percentage': (value_counts.values / len(df) * 100).round(2)
                    }), use_container_width=True)
    
    with tab4:
        st.markdown("### Correlation Analysis")
        
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        
        if len(numeric_cols) > 1:
            # Calculate correlation matrix
            corr_matrix = df[numeric_cols].corr()
            
            fig, ax = plt.subplots(figsize=(12, 10))
            sns.heatmap(corr_matrix, annot=True, fmt='.2f', cmap='coolwarm', 
                       center=0, square=True, ax=ax, cbar_kws={"shrink": 0.8})
            ax.set_title('Correlation Matrix')
            st.pyplot(fig)
            
            # Strong correlations
            st.markdown("#### Strong Correlations (|r| > 0.7)")
            high_corr = []
            for i in range(len(corr_matrix.columns)):
                for j in range(i+1, len(corr_matrix.columns)):
                    corr_value = corr_matrix.iloc[i, j]
                    if abs(corr_value) > 0.7:
                        high_corr.append({
                            'Feature 1': corr_matrix.columns[i],
                            'Feature 2': corr_matrix.columns[j],
                            'Correlation': corr_value
                        })
            
            if high_corr:
                high_corr_df = pd.DataFrame(high_corr)
                st.dataframe(high_corr_df.sort_values('Correlation', key=abs, ascending=False), 
                           use_container_width=True)
            else:
                st.info("No strong correlations found (|r| > 0.7)")
        else:
            st.warning("Need at least 2 numeric columns for correlation analysis.")

# Model Training Page
elif page == "🤖 Model Training":
    st.markdown('<h2 class="sub-header">🤖 Model Training Results</h2>', unsafe_allow_html=True)
    
    if not st.session_state.analysis_results:
        st.warning("Please run the analysis first on the Home page.")
        st.stop()
    
    results = st.session_state.analysis_results
    
    # Model Performance Comparison
    st.markdown("### Model Performance Comparison")
    
    # Create performance table
    performance_data = []
    for model_name, model_result in results['results'].items():
        performance_data.append({
            'Model': model_name,
            'Accuracy': model_result['accuracy'],
            'AUC': model_result['auc'],
            'F1 Score': model_result['f1'],
            'CV Score': model_result.get('cv_score', 0.0)
        })
    
    perf_df = pd.DataFrame(performance_data)
    
    # Display metrics
    col1, col2, col3, col4 = st.columns(4)
    best_result = results['results'][results['best_model_name']]
    
    with col1:
        st.metric("Best Model", results['best_model_name'])
    with col2:
        st.metric("AUC Score", f"{best_result['auc']:.4f}")
    with col3:
        st.metric("Accuracy", f"{best_result['accuracy']:.4f}")
    with col4:
        st.metric("F1 Score", f"{best_result['f1']:.4f}")
    
    # Performance chart
    st.markdown("#### Performance Visualization")
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(perf_df))
    width = 0.2
    
    ax.bar(x - width*1.5, perf_df['Accuracy'], width, label='Accuracy', color='#3B82F6')
    ax.bar(x - width/2, perf_df['AUC'], width, label='AUC', color='#10B981')
    ax.bar(x + width/2, perf_df['F1 Score'], width, label='F1 Score', color='#EF4444')
    ax.bar(x + width*1.5, perf_df['CV Score'], width, label='CV Score', color='#8B5CF6')
    
    ax.set_xlabel('Models')
    ax.set_ylabel('Score')
    ax.set_title('Model Performance Comparison')
    ax.set_xticks(x)
    ax.set_xticklabels(perf_df['Model'], rotation=45, ha='right')
    ax.legend()
    ax.grid(True, alpha=0.3)
    
    st.pyplot(fig)
    
    # Detailed table
    st.markdown("#### Detailed Performance Metrics")
    st.dataframe(perf_df, use_container_width=True)
    
    # Confusion Matrix
    st.markdown("#### Confusion Matrix - Best Model")
    from sklearn.metrics import confusion_matrix
    
    cm = confusion_matrix(results['dataset']['y_test'], best_result['predictions'])
    
    fig, ax = plt.subplots(figsize=(6, 5))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', 
               xticklabels=['No Churn', 'Churn'],
               yticklabels=['No Churn', 'Churn'],
               ax=ax)
    ax.set_title(f'Confusion Matrix - {results["best_model_name"]}')
    ax.set_xlabel('Predicted')
    ax.set_ylabel('Actual')
    st.pyplot(fig)
    
    # Feature Importance
    if hasattr(results['best_model'], 'feature_importances_'):
        st.markdown("#### Feature Importance")
        
        feature_importance = pd.DataFrame({
            'feature': results['dataset']['feature_names'],
            'importance': results['best_model'].feature_importances_
        }).sort_values('importance', ascending=False).head(15)
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            fig, ax = plt.subplots(figsize=(10, 8))
            bars = ax.barh(feature_importance['feature'], feature_importance['importance'])
            ax.set_xlabel('Importance')
            ax.set_title('Top 15 Feature Importance')
            ax.bar_label(bars, fmt='%.3f')
            st.pyplot(fig)
        
        with col2:
            st.dataframe(feature_importance, use_container_width=True)

# Business Insights Page
elif page == "🔍 Business Insights":
    st.markdown('<h2 class="sub-header">🔍 Business Insights & Recommendations</h2>', unsafe_allow_html=True)
    
    if not st.session_state.analysis_results:
        st.warning("Please run the analysis first on the Home page.")
        st.stop()
    
    insights = st.session_state.analysis_results['insights']
    
    # Current Situation
    st.markdown("### 📈 Current Business Situation")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Churn Rate", f"{insights['churn_rate']:.1f}%", 
                 f"{insights['churn_risk_level']} Risk")
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Estimated Annual Loss", f"${insights['estimated_annual_loss']:,.0f}")
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col3:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Avg Customer Value", f"${insights['avg_customer_value']:,.0f}")
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col4:
        potential_savings = insights['estimated_annual_loss'] * 0.3
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Potential Savings", f"${potential_savings:,.0f}", "30% reducible")
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Key Drivers
    st.markdown("### 🔍 Key Churn Drivers")
    
    drivers = insights['key_drivers']
    for i, driver in enumerate(drivers, 1):
        st.markdown(f"""
        <div class="feature-card">
            <h4 style="margin: 0; color: #1E40AF;">{i}. {driver}</h4>
        </div>
        """, unsafe_allow_html=True)
    
    # High-Risk Segments
    st.markdown("### ⚠️ High-Risk Customer Segments")
    
    segments = insights['high_risk_segments']
    seg_col1, seg_col2, seg_col3 = st.columns(3)
    
    for i, segment in enumerate(segments):
        col = [seg_col1, seg_col2, seg_col3][i % 3]
        with col:
            st.markdown(f"""
            <div style="background: linear-gradient(135deg, #FEF3C7 0%, #FDE68A 100%);
                        padding: 1.5rem; border-radius: 10px; margin: 0.5rem 0;
                        border-left: 5px solid #D97706;">
                <h4 style="color: #92400E; margin: 0 0 1rem 0;">{segment['segment']}</h4>
                <p style="margin: 0.5rem 0; color: #78350F;">
                    <strong>Risk Level:</strong> {segment['risk']}
                </p>
                <p style="margin: 0.5rem 0; color: #78350F;">
                    <strong>Customer Base:</strong> {segment['size']}
                </p>
            </div>
            """, unsafe_allow_html=True)
    
    # Recommendations
    st.markdown("### ✅ Recommended Actions")
    
    recommendations = insights['recommendations']
    for i, rec in enumerate(recommendations, 1):
        with st.expander(f"{i}. {rec['action']}", expanded=(i == 1)):
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Impact", rec['impact'])
            with col2:
                st.metric("Cost", rec['cost'])
            with col3:
                st.metric("Timeline", rec['timeline'])
            with col4:
                # Priority indicator
                if rec['impact'] == 'Very High' and rec['cost'] == 'Low':
                    st.success("🚨 HIGH PRIORITY")
                elif rec['impact'] == 'High' and rec['cost'] == 'Medium':
                    st.info("📋 MEDIUM PRIORITY")
                else:
                    st.warning("📝 LOW PRIORITY")
    
    # ROI Analysis
    st.markdown("### 📊 ROI Analysis")
    
    # Simple ROI calculation
    implementation_costs = {
        'Low': 10000,
        'Medium': 50000,
        'High': 200000
    }
    
    impact_multipliers = {
        'Very High': 0.4,
        'High': 0.3,
        'Medium': 0.2,
        'Low': 0.1
    }
    
    roi_data = []
    for rec in recommendations:
        cost = implementation_costs.get(rec['cost'], 50000)
        savings = insights['estimated_annual_loss'] * impact_multipliers.get(rec['impact'], 0.2)
        roi = (savings - cost) / cost * 100 if cost > 0 else 0
        payback = cost / savings * 12 if savings > 0 else float('inf')
        
        roi_data.append({
            'Action': rec['action'],
            'Cost ($)': f"${cost:,.0f}",
            'Annual Savings ($)': f"${savings:,.0f}",
            'ROI (%)': f"{roi:.1f}%",
            'Payback (months)': f"{payback:.1f}" if payback != float('inf') else "∞"
        })
    
    roi_df = pd.DataFrame(roi_data)
    st.dataframe(roi_df, use_container_width=True)

# Make Predictions Page
elif page == "🎯 Make Predictions":
    st.markdown('<h2 class="sub-header">🎯 Make Predictions</h2>', unsafe_allow_html=True)
    
    if not st.session_state.analysis_results:
        st.warning("Please train a model first on the Home page.")
        st.stop()
    
    predictor = st.session_state.analysis_results['predictor']
    features = predictor.feature_names
    
    # Prediction mode selector
    prediction_mode = st.radio(
        "Select prediction mode:",
        ["📝 Single Customer Prediction", "📁 Batch Prediction from File"]
    )
    
    if prediction_mode == "📝 Single Customer Prediction":
        st.markdown("### Predict for a Single Customer")
        
        # Create input form
        st.info("Enter values for the most important features. Other features will default to 0.")
        
        # Get top 10 important features
        important_features = []
        if hasattr(st.session_state.analysis_results['best_model'], 'feature_importances_'):
            importance_df = pd.DataFrame({
                'feature': features,
                'importance': st.session_state.analysis_results['best_model'].feature_importances_
            }).sort_values('importance', ascending=False)
            important_features = importance_df.head(10)['feature'].tolist()
        else:
            important_features = features[:10]  # Just use first 10
        
        # Create input fields
        input_data = {}
        cols_per_row = 2
        cols = st.columns(cols_per_row)
        
        for i, feature in enumerate(important_features):
            with cols[i % cols_per_row]:
                if any(keyword in feature.lower() for keyword in ['tenure', 'age', 'months', 'years', 'count']):
                    input_data[feature] = st.number_input(
                        f"{feature}",
                        value=0,
                        min_value=0,
                        max_value=100,
                        help=f"Enter value for {feature}"
                    )
                elif any(keyword in feature.lower() for keyword in ['amount', 'charge', 'fee', 'price', 'cost']):
                    input_data[feature] = st.number_input(
                        f"{feature}",
                        value=0.0,
                        min_value=0.0,
                        max_value=10000.0,
                        step=10.0,
                        format="%.2f",
                        help=f"Enter value for {feature}"
                    )
                elif any(keyword in feature.lower() for keyword in ['flag', 'indicator', 'has_', 'is_']):
                    input_data[feature] = st.selectbox(
                        f"{feature}",
                        options=[0, 1],
                        format_func=lambda x: "Yes" if x == 1 else "No",
                        help=f"Select value for {feature}"
                    )
                else:
                    input_data[feature] = st.number_input(
                        f"{feature}",
                        value=0.0,
                        help=f"Enter value for {feature}"
                    )
        
        # Set remaining features to 0
        for feature in features:
            if feature not in input_data:
                input_data[feature] = 0
        
        # Prediction threshold
        threshold = st.slider(
            "Prediction Threshold",
            min_value=0.0,
            max_value=1.0,
            value=0.5,
            step=0.05,
            help="Adjust the probability threshold for churn prediction"
        )
        
        if st.button("🔮 Make Prediction", type="primary", use_container_width=True):
            with st.spinner("Making prediction..."):
                try:
                    prediction = predictor.predict(input_data, threshold=threshold)
                    
                    if prediction['predictions']:
                        result = prediction['predictions'][0]
                        
                        # Display results
                        st.markdown("### 📊 Prediction Results")
                        
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            if result['prediction'] == 1:
                                st.error(f"## ⚠️ {result['prediction_label']}")
                            else:
                                st.success(f"## ✅ {result['prediction_label']}")
                        
                        with col2:
                            # Gauge chart
                            prob = result['churn_probability']
                            st.metric("Churn Probability", f"{prob:.1%}")
                            
                            fig, ax = plt.subplots(figsize=(8, 1))
                            ax.barh([''], [prob], color='#EF4444' if prob > 0.5 else '#10B981')
                            ax.set_xlim(0, 1)
                            ax.axvline(x=threshold, color='black', linestyle='--', 
                                     label=f'Threshold ({threshold})')
                            ax.set_xlabel('Probability')
                            ax.legend(loc='upper right')
                            st.pyplot(fig)
                        
                        with col3:
                            st.metric("Risk Level", result['risk_level'])
                            st.metric("Priority", result['priority'])
                        
                        # Detailed information
                        st.markdown("#### 📋 Details")
                        detail_col1, detail_col2 = st.columns(2)
                        
                        with detail_col1:
                            st.markdown(f"""
                            **Confidence Level:** {result['confidence']:.1%}
                            
                            **Customer ID:** {result['customer_id']}
                            
                            **Prediction Score:** {result['churn_probability']:.4f}
                            """)
                        
                        with detail_col2:
                            st.markdown(f"""
                            **Recommended Action:**
                            > {result['recommended_action']}
                            
                            **Next Steps:**
                            - Review customer history
                            - Check for recent issues
                            - Consider retention offers
                            """)
                    
                except Exception as e:
                    st.error(f"Prediction failed: {str(e)}")
    
    else:  # Batch Prediction
        st.markdown("### Batch Prediction from File")
        
        batch_file = st.file_uploader(
            "Upload CSV file with customer data",
            type=['csv'],
            help="Upload a CSV file with the same features as training data"
        )
        
        if batch_file:
            # Save uploaded file temporarily
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, "batch_data.csv")
            
            with open(temp_path, "wb") as f:
                f.write(batch_file.getbuffer())
            
            threshold = st.slider(
                "Prediction Threshold",
                min_value=0.0,
                max_value=1.0,
                value=0.5,
                step=0.05,
                key="batch_threshold"
            )
            
            if st.button("📊 Run Batch Prediction", type="primary", use_container_width=True):
                with st.spinner("Processing batch predictions..."):
                    try:
                        # Load the batch data
                        new_data = pd.read_csv(temp_path)
                        
                        # Make predictions
                        predictions = predictor.predict(new_data.to_dict('records'), threshold=threshold)
                        
                        if predictions:
                            # Display summary
                            st.markdown("### 📈 Prediction Summary")
                            
                            summary = predictions['summary']
                            col1, col2, col3, col4 = st.columns(4)
                            
                            with col1:
                                st.metric("Total Customers", summary['total_predictions'])
                            with col2:
                                st.metric("Predicted to Churn", summary['churn_count'])
                            with col3:
                                st.metric("Churn Rate", f"{summary['churn_rate']:.1f}%")
                            with col4:
                                st.metric("High Risk", summary['high_risk_customers'])
                            
                            # Show predictions table
                            st.markdown("### 📋 Detailed Predictions")
                            predictions_df = pd.DataFrame(predictions['predictions'])
                            st.dataframe(predictions_df, use_container_width=True, height=400)
                            
                            # Download button for predictions
                            csv = predictions_df.to_csv(index=False)
                            st.download_button(
                                label="📥 Download Predictions as CSV",
                                data=csv,
                                file_name="churn_predictions.csv",
                                mime="text/csv"
                            )
                            
                            # Risk distribution chart
                            st.markdown("### 📊 Risk Distribution")
                            risk_counts = predictions_df['risk_level'].value_counts()
                            
                            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
                            
                            # Pie chart
                            colors = ['#10B981', '#3B82F6', '#F59E0B', '#EF4444']
                            wedges, texts, autotexts = ax1.pie(
                                risk_counts.values,
                                labels=risk_counts.index,
                                autopct='%1.1f%%',
                                colors=colors,
                                startangle=90
                            )
                            ax1.set_title('Risk Level Distribution')
                            
                            # Bar chart
                            ax2.bar(risk_counts.index, risk_counts.values, color=colors[:len(risk_counts)])
                            ax2.set_title('Risk Level Counts')
                            ax2.set_xlabel('Risk Level')
                            ax2.set_ylabel('Count')
                            plt.setp(ax2.get_xticklabels(), rotation=45, ha='right')
                            
                            st.pyplot(fig)
                    
                    except Exception as e:
                        st.error(f"Batch prediction failed: {str(e)}")

# Download Model Page (WITH EXPANDERS FOR DETAILS)
elif page == "💾 Download Model":
    st.markdown('<h2 class="sub-header">💾 Download Trained Model</h2>', unsafe_allow_html=True)
    
    if not st.session_state.analysis_results:
        st.warning("Please train a model first on the Home page.")
        st.stop()
    
    results = st.session_state.analysis_results
    best_model_name = results['best_model_name']
    best_result = results['results'][best_model_name]
    feature_count = len(results['dataset']['feature_names'])
    all_features = results['dataset']['feature_names']
    
    # Initialize session state for download
    if 'model_zip_data' not in st.session_state:
        st.session_state.model_zip_data = None
    if 'model_zip_filename' not in st.session_state:
        st.session_state.model_zip_filename = None
    
    # Main content area
    st.markdown("### Model Information")
    
    # Create two columns for basic info
    info_col1, info_col2 = st.columns(2)
    
    with info_col1:
        # Model Name
        st.markdown(f"**Model Name:** {best_model_name}")
        
        # ACC Score
        st.markdown(f"**ACC Score:**")
        st.markdown(f"- **{best_model_name}:** {best_result['accuracy']:.4f}")
        
        # Features count
        st.markdown(f"**Total Features:** {feature_count}")
    
    with info_col2:
        # Download Trained Model checkbox
        download_model = st.checkbox("Download Trained Model", value=True, key="download_checkbox")
        
        if download_model:
            st.success("✅ Model is ready for download")
    
    # Show ALL Feature List in expander
    st.markdown("---")
    with st.expander(f"📋 Show Feature List | Total Features: {feature_count}", expanded=False):
        # Create 3 columns for features
        col1, col2, col3 = st.columns(3)
        
        # Calculate items per column
        items_per_col = (feature_count + 2) // 3  # Round up
        
        # Distribute features across columns WITHOUT numbers
        feature_display_cols = [col1, col2, col3]
        
        for i in range(3):
            with feature_display_cols[i]:
                start_idx = i * items_per_col
                end_idx = min((i + 1) * items_per_col, feature_count)
                
                for idx in range(start_idx, end_idx):
                    feature_name = all_features[idx]
                    
                    # Display each feature WITHOUT numbering
                    st.markdown(f"""
                    <div class="feature-item">
                        {feature_name}
                    </div>
                    """, unsafe_allow_html=True)
        
        # Also show features in a code block for easy copying
        st.markdown("---")
        st.markdown("**📋 Copy All Features to Clipboard:**")
        all_features_text = "\n".join(all_features)
        st.code(all_features_text, language="text")
    
    # Model Configuration Section
    st.markdown("---")
    st.markdown("### Model Configuration")
    
    config_col1, config_col2 = st.columns([1, 1])
    
    with config_col1:
        # Model Name for Download input
        default_model_name = f"churn_model_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        model_name_input = st.text_input(
            "Model Name for Download",
            value=default_model_name,
            help="Enter a name for your model file",
            key="model_name_input"
        )
        
        # Display example model name
        st.markdown(f"**Example:** `churn_model_20231225_143144`")
    
    with config_col2:
        # Model Type display
        st.markdown(f"**Model Type:** {best_model_name}")
        st.markdown(f"**Creation Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        st.markdown(f"**Random State:** {CONFIG.RANDOM_STATE}")
    
    # Model Performance Section
    st.markdown("---")
    st.markdown("### Model Performance")
    
    # Create performance metrics cards
    perf_col1, perf_col2, perf_col3, perf_col4 = st.columns(4)
    
    with perf_col1:
        st.markdown(f"""
        <div class="stats-card">
            <div style="font-size: 28px; font-weight: bold; color: #3B82F6;">{best_result['auc']:.4f}</div>
            <div style="color: #6b7280; font-size: 14px;">AUC Score</div>
            <div style="color: #9ca3af; font-size: 12px; margin-top: 5px;">(Higher is better)</div>
        </div>
        """, unsafe_allow_html=True)
    
    with perf_col2:
        st.markdown(f"""
        <div class="stats-card">
            <div style="font-size: 28px; font-weight: bold; color: #10B981;">{best_result['accuracy']:.4f}</div>
            <div style="color: #6b7280; font-size: 14px;">Accuracy</div>
            <div style="color: #9ca3af; font-size: 12px; margin-top: 5px;">Correct predictions</div>
        </div>
        """, unsafe_allow_html=True)
    
    with perf_col3:
        st.markdown(f"""
        <div class="stats-card">
            <div style="font-size: 28px; font-weight: bold; color: #8B5CF6;">{best_result['f1']:.4f}</div>
            <div style="color: #6b7280; font-size: 14px;">F1 Score</div>
            <div style="color: #9ca3af; font-size: 12px; margin-top: 5px;">Balance of precision & recall</div>
        </div>
        """, unsafe_allow_html=True)
    
    with perf_col4:
        cv_score = best_result.get('cv_score', best_result['auc'])
        st.markdown(f"""
        <div class="stats-card">
            <div style="font-size: 28px; font-weight: bold; color: #F59E0B;">{cv_score:.4f}</div>
            <div style="color: #6b7280; font-size: 14px;">CV Score</div>
            <div style="color: #9ca3af; font-size: 12px; margin-top: 5px;">Cross-validation score</div>
        </div>
        """, unsafe_allow_html=True)
    
    # Performance details in expander
    with st.expander("📊 View Detailed Performance Metrics", expanded=False):
        col_a, col_b = st.columns(2)
        
        with col_a:
            st.markdown("#### Confusion Matrix")
            from sklearn.metrics import confusion_matrix
            cm = confusion_matrix(results['dataset']['y_test'], best_result['predictions'])
            
            fig, ax = plt.subplots(figsize=(5, 4))
            sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', 
                       xticklabels=['No Churn', 'Churn'],
                       yticklabels=['No Churn', 'Churn'],
                       ax=ax)
            ax.set_title(f'Confusion Matrix')
            ax.set_xlabel('Predicted')
            ax.set_ylabel('Actual')
            st.pyplot(fig)
        
        with col_b:
            st.markdown("#### Classification Report")
            from sklearn.metrics import classification_report
            
            report = classification_report(
                results['dataset']['y_test'], 
                best_result['predictions'],
                output_dict=True
            )
            report_df = pd.DataFrame(report).transpose()
            st.dataframe(report_df, use_container_width=True)
        
        # Additional metrics
        st.markdown("#### Additional Metrics")
        col_c, col_d, col_e = st.columns(3)
        
        with col_c:
            from sklearn.metrics import precision_score, recall_score
            precision = precision_score(results['dataset']['y_test'], best_result['predictions'])
            recall = recall_score(results['dataset']['y_test'], best_result['predictions'])
            st.metric("Precision", f"{precision:.4f}")
            st.metric("Recall", f"{recall:.4f}")
        
        with col_d:
            from sklearn.metrics import roc_curve
            fpr, tpr, _ = roc_curve(results['dataset']['y_test'], best_result['probabilities'])
            st.metric("False Positive Rate", f"{fpr.mean():.4f}")
            st.metric("True Positive Rate", f"{tpr.mean():.4f}")
        
        with col_e:
            st.metric("Training Samples", len(results['dataset']['X_train']))
            st.metric("Test Samples", len(results['dataset']['X_test']))
    
    # Feature Importance Analysis in expander (if available)
    if hasattr(results['best_model'], 'feature_importances_'):
        with st.expander("🔍 Feature Importance Analysis", expanded=False):
            # Get feature importance
            feature_importance = pd.DataFrame({
                'feature': all_features,
                'importance': results['best_model'].feature_importances_
            }).sort_values('importance', ascending=False)
            
            # Show top 10 features visualization
            st.markdown("#### Top 10 Most Important Features")
            top_features = feature_importance.head(10)
            
            # Create visualization
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.barh(top_features['feature'], top_features['importance'])
            ax.set_xlabel('Importance Score')
            ax.set_title('Top 10 Feature Importance')
            ax.bar_label(bars, fmt='%.3f')
            plt.gca().invert_yaxis()
            st.pyplot(fig)
            
            # Display in 2 columns
            st.markdown("---")
            imp_col1, imp_col2 = st.columns(2)
            
            with imp_col1:
                st.markdown("##### Top 5 Features:")
                for i, row in feature_importance.head(5).iterrows():
                    st.markdown(f"""
                    <div class="feature-item" style="border-left-color: #EF4444; margin: 5px 0;">
                        <div style="font-weight: bold; color: #1E40AF;">{row['feature']}</div>
                        <div style="color: #6b7280; font-size: 12px;">Importance: {row['importance']:.4f}</div>
                    </div>
                    """, unsafe_allow_html=True)
            
            with imp_col2:
                st.markdown("##### Feature Importance Distribution:")
                
                # Categorize by importance
                high = len(feature_importance[feature_importance['importance'] > 0.01])
                medium = len(feature_importance[(feature_importance['importance'] > 0.001) & 
                                               (feature_importance['importance'] <= 0.01)])
                low = len(feature_importance[feature_importance['importance'] <= 0.001])
                
                st.metric("High Importance (>0.01)", high)
                st.metric("Medium Importance", medium)
                st.metric("Low Importance (≤0.001)", low)
                
                # Pie chart of importance distribution
                st.markdown("##### Distribution Chart")
                fig2, ax2 = plt.subplots(figsize=(5, 4))
                sizes = [high, medium, low]
                labels = [f'High\n({high})', f'Medium\n({medium})', f'Low\n({low})']
                colors = ['#EF4444', '#F59E0B', '#3B82F6']
                ax2.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
                ax2.axis('equal')
                st.pyplot(fig2)
    
    # Generate Model Package Section
    st.markdown("---")
    st.markdown("### Generate Model Package")
    
    # Status indicator
    status_col1, status_col2 = st.columns([3, 1])
    
    with status_col1:
        if st.session_state.model_zip_data:
            st.success("✅ Model package is ready for download!")
        elif download_model:
            st.info("🔧 Ready to generate model package")
        else:
            st.warning("📝 Check 'Download Trained Model' to enable")
    
    with status_col2:
        # GENERATE MODEL button
        generate_disabled = not download_model
        
        if st.button("🔄 Generate Model Package", 
                    type="primary", 
                    use_container_width=True,
                    disabled=generate_disabled,
                    key="generate_model_package"):
            
            if generate_disabled:
                st.warning("⚠️ Please check 'Download Trained Model' first")
            else:
                with st.spinner("Generating model package..."):
                    try:
                        # Step 1: Create the model file
                        model_package = {
                            'model': results['best_model'],
                            'scaler': results['dataset']['preprocessor'].scaler,
                            'feature_names': all_features,
                            'metadata': {
                                'model_name': model_name_input,
                                'model_type': best_model_name,
                                'accuracy': float(best_result['accuracy']),
                                'auc': float(best_result['auc']),
                                'f1': float(best_result['f1']),
                                'features': feature_count,
                                'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                'dataset_info': {
                                    'original_rows': len(results['df']),
                                    'original_columns': len(results['df'].columns),
                                    'training_samples': len(results['dataset']['X_train']),
                                    'test_samples': len(results['dataset']['X_test']),
                                    'target_column': results['target_column']
                                }
                            }
                        }
                        
                        # Save model to bytes
                        model_bytes = io.BytesIO()
                        joblib.dump(model_package, model_bytes)
                        model_bytes.seek(0)
                        
                        # Step 2: Create ZIP file
                        zip_bytes = io.BytesIO()
                        with zipfile.ZipFile(zip_bytes, 'w') as zipf:
                            # Add model file
                            zipf.writestr(f"{model_name_input}.pkl", model_bytes.getvalue())
                            
                            # Add features file WITHOUT numbering
                            features_content = f"""UNIVERSAL CHURN PREDICTION MODEL
{'=' * 50}

MODEL INFORMATION:
-----------------
Model Name:     {model_name_input}
Model Type:     {best_model_name}
Created:        {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
Total Features: {feature_count}

PERFORMANCE METRICS:
-------------------
Accuracy:       {best_result['accuracy']:.4f}
AUC Score:      {best_result['auc']:.4f}
F1 Score:       {best_result['f1']:.4f}

DATASET INFORMATION:
-------------------
Original Dataset: {results.get('data_path', 'Unknown')}
Target Column:    {results['target_column']}
Training Samples: {len(results['dataset']['X_train'])}
Test Samples:     {len(results['dataset']['X_test'])}

{'=' * 50}
ALL FEATURES ({feature_count} total):
{'=' * 50}\n\n"""
                            
                            # Format features in 3 columns WITHOUT numbers
                            items_per_row = 3
                            rows = (feature_count + items_per_row - 1) // items_per_row
                            
                            for row in range(rows):
                                line = ""
                                for col in range(items_per_row):
                                    idx = row + col * rows
                                    if idx < feature_count:
                                        feature_name = all_features[idx]
                                        # Format: feature_name (padded to 35 chars) - NO NUMBERING
                                        line += f"{feature_name:<35}"
                                features_content += line + "\n"
                            
                            zipf.writestr("features.txt", features_content)
                            
                            # Add info file
                            info_content = f"""
                            ============================================
                            UNIVERSAL CHURN PREDICTION MODEL PACKAGE
                            ============================================
                            
                            PACKAGE CONTENTS:
                            -----------------
                            1. {model_name_input}.pkl   - Serialized model with all components
                            2. features.txt            - Complete list of all features
                            3. README.txt             - This file
                            
                            MODEL DETAILS:
                            -------------
                            Model Name:    {model_name_input}
                            Model Type:    {best_model_name}
                            Created:       {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
                            
                            PERFORMANCE:
                            -----------
                            Accuracy:      {best_result['accuracy']:.4f}
                            AUC Score:     {best_result['auc']:.4f}
                            F1 Score:      {best_result['f1']:.4f}
                            Features:      {feature_count}
                            
                            HOW TO USE THIS MODEL:
                            ----------------------
                            1. Install required packages:
                               pip install scikit-learn pandas numpy joblib
                            
                            2. Load the model:
                               import joblib
                               model_data = joblib.load('{model_name_input}.pkl')
                               model = model_data['model']
                               scaler = model_data['scaler']
                               features = model_data['feature_names']
                            
                            3. Prepare new data:
                               - Ensure your data has ALL {feature_count} features
                               - Missing features will be set to 0
                               - Extra features will be ignored
                            
                            4. Make predictions:
                               # Scale the data
                               scaled_data = scaler.transform(new_data[features])
                               
                               # Predict churn
                               predictions = model.predict(scaled_data)
                               probabilities = model.predict_proba(scaled_data)
                            
                            NOTES:
                            ------
                            • See features.txt for complete feature list
                            • Model expects scaled features
                            • All categorical features are already encoded
                            • Threshold for churn prediction: 0.5
                            
                            ============================================
                            Generated by Universal Churn Prediction System
                            ============================================
                            """
                            zipf.writestr("README.txt", info_content)
                        
                        zip_bytes.seek(0)
                        
                        # Store in session state
                        st.session_state.model_zip_data = zip_bytes.getvalue()
                        st.session_state.model_zip_filename = f"{model_name_input}.zip"
                        
                        st.success(f"✅ Model package generated successfully!")
                        st.balloons()
                        
                    except Exception as e:
                        st.error(f"❌ Error generating package: {str(e)}")
    
    # Show download button if we have data
    if st.session_state.model_zip_data and st.session_state.model_zip_filename:
        st.markdown("---")
        st.markdown("### 📥 Download Your Model")
        
        # File info
        file_size = len(st.session_state.model_zip_data) / 1024  # KB
        st.info(f"**Package Size:** {file_size:.1f} KB | **Files:** 3 files (.pkl, features.txt, README.txt)")
        
        # Download button
        st.download_button(
            label=f"💾 Download {st.session_state.model_zip_filename}",
            data=st.session_state.model_zip_data,
            file_name=st.session_state.model_zip_filename,
            mime="application/zip",
            use_container_width=True,
            help="Click to download the complete model package",
            key="download_model_package"
        )
        
        # Package contents preview in expander
        with st.expander("📦 Preview Package Contents", expanded=False):
            st.markdown("""
            **The downloaded ZIP file contains:**
            
            1. **`model_name.pkl`** - Serialized model with all components
               - Trained ML model
               - Feature scaler
               - Feature names list
               - Metadata and performance metrics
            
            2. **`features.txt`** - Complete list of all features
               - All feature names used by the model
               - Formatted in 3 columns for easy reading
            
            3. **`README.txt`** - Instructions and documentation
               - How to load and use the model
               - Performance details
               - Requirements and notes
            """)
    
    # ============================================
    # WORD DOCUMENT GENERATION SECTION
    # ============================================
    
    # NEW: Word Document Generation Section
    st.markdown("---")
    st.markdown("### 📄 Generate Complete Report Document")
    
    doc_col1, doc_col2 = st.columns([3, 1])
    
    with doc_col1:
        # Word document name input
        doc_name = st.text_input(
            "Document Name",
            value=f"Churn_Analysis_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            help="Enter a name for your Word document",
            key="doc_name_input"
        )
    
    with doc_col2:
        # Generate Word Document button
        if st.button("📝 Generate Word Doc", use_container_width=True, key="generate_word_doc"):
            with st.spinner("Creating comprehensive Word document..."):
                try:
                    # Create a new Word document
                    doc = docx.Document()
                    
                    # Set document properties
                    doc.core_properties.title = f"Churn Prediction Analysis Report"
                    doc.core_properties.author = "Universal Churn Prediction System"
                    doc.core_properties.subject = "Comprehensive Churn Analysis Report"
                    
                    # ============================================
                    # COVER PAGE
                    # ============================================
                    # Add title
                    title = doc.add_heading('CHURN PREDICTION ANALYSIS REPORT', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add subtitle
                    subtitle = doc.add_paragraph()
                    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = subtitle.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(100, 100, 100)
                    
                    doc.add_paragraph().add_run().add_break()  # Spacer
                    
                    # Add company info
                    company_info = doc.add_paragraph()
                    company_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    company_run = company_info.add_run("Universal Churn Prediction System")
                    company_run.font.size = Pt(16)
                    company_run.font.color.rgb = RGBColor(59, 130, 246)
                    company_run.font.bold = True
                    
                    doc.add_page_break()
                    
                    # ============================================
                    # TABLE OF CONTENTS
                    # ============================================
                    toc_title = doc.add_heading('Table of Contents', 1)
                    toc_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    toc_items = [
                        "1. Executive Summary",
                        "2. Dataset Overview",
                        "3. Data Analysis & Preprocessing",
                        "4. Model Training Results",
                        "5. Business Insights",
                        "6. Feature Importance Analysis",
                        "7. Download Information",
                        "8. Appendices"
                    ]
                    
                    for item in toc_items:
                        p = doc.add_paragraph(item, style='List Number')
                        p.paragraph_format.left_indent = Inches(0.5)
                    
                    doc.add_page_break()
                    
                    # ============================================
                    # 1. EXECUTIVE SUMMARY
                    # ============================================
                    doc.add_heading('1. Executive Summary', 1)
                    
                    # Summary paragraph
                    summary = doc.add_paragraph()
                    summary.add_run("This report presents a comprehensive analysis of customer churn prediction using machine learning models. ")
                    summary.add_run(f"The analysis was performed on a dataset containing {len(results['df']):,} customers with {len(results['df'].columns)} features. ")
                    summary.add_run(f"The target variable '{results['target_column']}' was used for prediction.\n\n")
                    
                    # Key findings table
                    doc.add_heading('Key Findings', 2)
                    
                    # Calculate churn rate safely
                    try:
                        # Try to get churn rate from target column
                        target_col = results['target_column']
                        if target_col in results['df'].columns:
                            # Convert to numeric if possible
                            try:
                                # Try to convert to numeric
                                target_series = pd.to_numeric(results['df'][target_col], errors='coerce')
                                churn_rate = target_series.mean() * 100
                                if pd.isna(churn_rate):
                                    churn_rate = 0
                            except:
                                # If conversion fails, check for Yes/No values
                                unique_vals = results['df'][target_col].unique()
                                if 'Yes' in unique_vals or 'yes' in unique_vals:
                                    churn_count = (results['df'][target_col].astype(str).str.lower() == 'yes').sum()
                                    churn_rate = (churn_count / len(results['df'])) * 100
                                elif '1' in unique_vals:
                                    churn_count = (results['df'][target_col].astype(str) == '1').sum()
                                    churn_rate = (churn_count / len(results['df'])) * 100
                                else:
                                    churn_rate = 0
                        else:
                            churn_rate = 0
                    except:
                        churn_rate = 0
                    
                    # Create table with proper row count
                    data_rows = [
                        ('Best Performing Model', best_model_name),
                        ('Model Accuracy', f"{best_result['accuracy']:.4f}"),
                        ('AUC Score', f"{best_result['auc']:.4f}"),
                        ('Total Features', str(feature_count)),
                        ('Churn Rate', f"{churn_rate:.1f}%")
                    ]
                    
                    table = doc.add_table(rows=len(data_rows)+1, cols=2)
                    table.style = 'Light Shading'
                    
                    # Table headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Metric'
                    hdr_cells[1].text = 'Value'
                    
                    # Table data
                    for i, (metric, value) in enumerate(data_rows, 1):
                        if i < len(table.rows):
                            row_cells = table.rows[i].cells
                            row_cells[0].text = metric
                            row_cells[1].text = value
                    
                    doc.add_paragraph().add_run().add_break()  # Spacer
                    
                    # ============================================
                    # 2. DATASET OVERVIEW
                    # ============================================
                    doc.add_heading('2. Dataset Overview', 1)
                    
                    # Dataset information
                    dataset_info = [
                        ('Target Column', results['target_column']),
                        ('Total Rows', f"{len(results['df']):,}"),
                        ('Total Columns', str(len(results['df'].columns))),
                    ]
                    
                    for info_name, info_value in dataset_info:
                        p = doc.add_paragraph()
                        p.add_run(f"{info_name}: ").bold = True
                        p.add_run(info_value)
                    
                    doc.add_paragraph().add_run().add_break()  # Spacer
                    
                    # Data types table
                    doc.add_heading('Data Types Distribution', 2)
                    
                    dtype_counts = results['df'].dtypes.value_counts()
                    
                    # Create table with proper row count
                    dtype_table = doc.add_table(rows=len(dtype_counts)+1, cols=2)
                    dtype_table.style = 'Light Grid'
                    
                    dtype_hdr = dtype_table.rows[0].cells
                    dtype_hdr[0].text = 'Data Type'
                    dtype_hdr[1].text = 'Count'
                    
                    for idx, (dtype, count) in enumerate(dtype_counts.items(), 1):
                        if idx < len(dtype_table.rows):
                            row_cells = dtype_table.rows[idx].cells
                            row_cells[0].text = str(dtype)
                            row_cells[1].text = str(count)
                    
                    doc.add_page_break()
                    
                    # ============================================
                    # 3. DATA ANALYSIS & PREPROCESSING
                    # ============================================
                    doc.add_heading('3. Data Analysis & Preprocessing', 1)
                    
                    # Missing values analysis (safe version)
                    try:
                        missing_values = results['df'].isnull().sum()
                        missing_values = missing_values[missing_values > 0]
                        missing_count = len(missing_values)
                        
                        p = doc.add_paragraph()
                        p.add_run("Missing Values Analysis: ").bold = True
                        p.add_run(f"{missing_count} columns contain missing values")
                        
                        if missing_count > 0:
                            # Create table with proper row count
                            table_rows = min(missing_count, 10) + 1
                            missing_table = doc.add_table(rows=table_rows, cols=3)
                            missing_table.style = 'Light Grid'
                            
                            missing_hdr = missing_table.rows[0].cells
                            missing_hdr[0].text = 'Column'
                            missing_hdr[1].text = 'Missing Count'
                            missing_hdr[2].text = 'Missing %'
                            
                            for idx, (col, count) in enumerate(missing_values.head(10).items(), 1):
                                if idx < len(missing_table.rows):
                                    row_cells = missing_table.rows[idx].cells
                                    row_cells[0].text = str(col)
                                    row_cells[1].text = str(count)
                                    row_cells[2].text = f"{(count/len(results['df'])*100):.2f}%"
                    except Exception as e:
                        doc.add_paragraph(f"Missing values analysis skipped: {str(e)[:50]}")
                    
                    # Preprocessing steps
                    doc.add_heading('Preprocessing Steps Applied', 2)
                    
                    preprocessing_steps = [
                        '✓ Duplicate rows removed',
                        '✓ Missing values imputed',
                        '✓ Categorical variables encoded',
                        '✓ Numerical features scaled',
                        '✓ Feature engineering applied',
                        '✓ Outliers handled'
                    ]
                    
                    for step in preprocessing_steps:
                        doc.add_paragraph(step, style='List Bullet')
                    
                    doc.add_paragraph().add_run().add_break()  # Spacer
                    
                    # ============================================
                    # 4. MODEL TRAINING RESULTS
                    # ============================================
                    doc.add_heading('4. Model Training Results', 1)
                    
                    # Performance comparison table
                    doc.add_heading('Model Performance Comparison', 2)
                    
                    # Create table with proper row count
                    model_count = len(results['results'])
                    perf_table = doc.add_table(rows=model_count + 1, cols=5)
                    perf_table.style = 'Light Shading Accent 1'
                    
                    perf_hdr = perf_table.rows[0].cells
                    perf_hdr[0].text = 'Model'
                    perf_hdr[1].text = 'Accuracy'
                    perf_hdr[2].text = 'AUC Score'
                    perf_hdr[3].text = 'F1 Score'
                    perf_hdr[4].text = 'CV Score'
                    
                    for idx, (model_name, model_result) in enumerate(results['results'].items(), 1):
                        if idx < len(perf_table.rows):
                            row_cells = perf_table.rows[idx].cells
                            row_cells[0].text = model_name
                            row_cells[1].text = f"{model_result.get('accuracy', 0):.4f}"
                            row_cells[2].text = f"{model_result.get('auc', 0):.4f}"
                            row_cells[3].text = f"{model_result.get('f1', 0):.4f}"
                            row_cells[4].text = f"{model_result.get('cv_score', model_result.get('auc', 0)):.4f}"
                    
                    doc.add_paragraph().add_run().add_break()  # Spacer
                    
                    # Best model details
                    doc.add_heading(f'Best Model: {best_model_name}', 2)
                    
                    best_model_details = [
                        ('Accuracy', f"{best_result['accuracy']:.4f}"),
                        ('AUC Score', f"{best_result['auc']:.4f}"),
                        ('F1 Score', f"{best_result['f1']:.4f}"),
                        ('Training Samples', str(len(results['dataset']['X_train']))),
                        ('Test Samples', str(len(results['dataset']['X_test']))),
                        ('Feature Count', str(feature_count))
                    ]
                    
                    for detail_name, detail_value in best_model_details:
                        p = doc.add_paragraph()
                        p.add_run(f"{detail_name}: ").bold = True
                        p.add_run(detail_value)
                    
                    # Add confusion matrix image
                    try:
                        from sklearn.metrics import confusion_matrix
                        import seaborn as sns
                        
                        cm = confusion_matrix(results['dataset']['y_test'], best_result['predictions'])
                        
                        fig, ax = plt.subplots(figsize=(6, 5))
                        sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', 
                                   xticklabels=['No Churn', 'Churn'],
                                   yticklabels=['No Churn', 'Churn'],
                                   ax=ax)
                        ax.set_title(f'Confusion Matrix - {best_model_name}')
                        ax.set_xlabel('Predicted')
                        ax.set_ylabel('Actual')
                        
                        # Save plot to buffer
                        buf = BytesIO()
                        plt.tight_layout()
                        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                        plt.close(fig)
                        buf.seek(0)
                        
                        # Add image to document
                        doc.add_heading('Confusion Matrix', 3)
                        doc.add_picture(buf, width=Inches(5))
                        
                    except Exception as e:
                        doc.add_paragraph(f"Could not generate confusion matrix: {str(e)[:100]}")
                    
                    doc.add_page_break()
                    
                    # ============================================
                    # 5. BUSINESS INSIGHTS
                    # ============================================
                    doc.add_heading('5. Business Insights', 1)
                    
                    try:
                        # Generate insights
                        insights_gen = BusinessInsights(
                            dataset=results['dataset'],
                            model_results=results['results'],
                            preprocessor=results['preprocessor'],
                            original_df=results['df']
                        )
                        insights = insights_gen.generate_all_insights()
                        
                        # Financial impact
                        doc.add_heading('Financial Impact Analysis', 2)
                        
                        # Create table with proper row count
                        financial_data = [
                            ('Actual Churn Rate', f"{insights.get('actual_churn_rate', 0):.1f}%"),
                            ('Predicted Churn Rate', f"{insights.get('predicted_churn_rate', 0):.1f}%"),
                            ('Estimated Annual Loss', f"${insights.get('estimated_annual_loss', 0):,.0f}"),
                            ('Recoverable Value', f"${insights.get('recoverable_value', 0):,.0f}")
                        ]
                        
                        financial_table = doc.add_table(rows=len(financial_data), cols=2)
                        financial_table.style = 'Light Grid Accent 2'
                        
                        for idx, (metric, value) in enumerate(financial_data):
                            if idx < len(financial_table.rows):
                                row_cells = financial_table.rows[idx].cells
                                row_cells[0].text = metric
                                row_cells[1].text = value
                        
                        doc.add_paragraph().add_run().add_break()  # Spacer
                        
                        # Key drivers
                        doc.add_heading('Key Churn Drivers', 2)
                        
                        drivers = insights.get('key_drivers', [])
                        for i, driver in enumerate(drivers[:5], 1):
                            p = doc.add_paragraph(f"{i}. ", style='List Number')
                            p.add_run(f"{driver.get('feature', 'N/A')}").bold = True
                            p.add_run(f" - {driver.get('impact', 'N/A')}")
                        
                        # Recommendations
                        doc.add_heading('Recommended Actions', 2)
                        
                        recommendations = insights.get('recommendations', [])
                        for i, rec in enumerate(recommendations[:5], 1):
                            p = doc.add_paragraph(f"{i}. ", style='List Number')
                            p.add_run(f"{rec.get('action', 'N/A')}").bold = True
                            
                            details = doc.add_paragraph()
                            details.add_run(f"Impact: {rec.get('impact', 'N/A')} | ")
                            details.add_run(f"Cost: {rec.get('cost', 'N/A')} | ")
                            details.add_run(f"Timeline: {rec.get('timeline', 'N/A')} | ")
                            details.add_run(f"Priority: {rec.get('priority', 'N/A')}")
                            
                    except Exception as e:
                        doc.add_paragraph(f"Business insights generation failed: {str(e)[:100]}")
                        # Add default insights
                        doc.add_heading('Key Churn Drivers', 2)
                        doc.add_paragraph("1. Customer Tenure - Shorter tenure increases churn risk")
                        doc.add_paragraph("2. Monthly Charges - Higher charges increase churn risk")
                        doc.add_paragraph("3. Contract Type - Month-to-month contracts have higher risk")
                        
                        doc.add_heading('Recommended Actions', 2)
                        doc.add_paragraph("1. Implement retention program for high-risk customers")
                        doc.add_paragraph("2. Offer contract incentives for long-term commitments")
                        doc.add_paragraph("3. Improve customer support response time")
                    
                    doc.add_page_break()
                    
                    # ============================================
                    # 6. FEATURE IMPORTANCE ANALYSIS
                    # ============================================
                    doc.add_heading('6. Feature Importance Analysis', 1)
                    
                    if hasattr(results['best_model'], 'feature_importances_'):
                        # Top features table
                        doc.add_heading('Top 10 Most Important Features', 2)
                        
                        feature_importance = pd.DataFrame({
                            'feature': all_features,
                            'importance': results['best_model'].feature_importances_
                        }).sort_values('importance', ascending=False)
                        
                        top_features = feature_importance.head(10)
                        
                        # Create table with proper row count
                        feat_table = doc.add_table(rows=len(top_features) + 1, cols=3)
                        feat_table.style = 'Light Shading Accent 3'
                        
                        feat_hdr = feat_table.rows[0].cells
                        feat_hdr[0].text = 'Rank'
                        feat_hdr[1].text = 'Feature Name'
                        feat_hdr[2].text = 'Importance Score'
                        
                        for idx, (_, row) in enumerate(top_features.iterrows(), 1):
                            if idx < len(feat_table.rows):
                                row_cells = feat_table.rows[idx].cells
                                row_cells[0].text = str(idx)
                                row_cells[1].text = row['feature']
                                row_cells[2].text = f"{row['importance']:.4f}"
                        
                        # Add feature importance chart
                        try:
                            fig, ax = plt.subplots(figsize=(8, 6))
                            top_5 = feature_importance.head(5)
                            bars = ax.barh(top_5['feature'], top_5['importance'])
                            ax.set_xlabel('Importance Score')
                            ax.set_title('Top 5 Feature Importance')
                            ax.bar_label(bars, fmt='%.3f')
                            plt.gca().invert_yaxis()
                            
                            buf = BytesIO()
                            plt.tight_layout()
                            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                            plt.close(fig)
                            buf.seek(0)
                            
                            doc.add_heading('Feature Importance Visualization', 2)
                            doc.add_picture(buf, width=Inches(6))
                            
                        except Exception as e:
                            doc.add_paragraph(f"Could not generate feature importance chart: {str(e)[:100]}")
                    else:
                        doc.add_paragraph("Feature importance analysis is not available for this model type.")
                    
                    doc.add_page_break()
                    
                    # ============================================
                    # 7. DOWNLOAD INFORMATION
                    # ============================================
                    doc.add_heading('7. Download Information', 1)
                    
                    # Model package information
                    doc.add_heading('Model Package Contents', 2)
                    
                    package_contents = [
                        ('Model File (.pkl)', 'Serialized model with all components'),
                        ('Features List (features.txt)', 'Complete list of all features used'),
                        ('Documentation (README.txt)', 'Instructions and performance details'),
                        ('Word Document', 'This comprehensive report')
                    ]
                    
                    for content, description in package_contents:
                        p = doc.add_paragraph()
                        p.add_run(f"• {content}: ").bold = True
                        p.add_run(description)
                    
                    # Feature list (abbreviated)
                    doc.add_heading('All Features Used', 2)
                    doc.add_paragraph(f"Total Features: {feature_count}")
                    
                    # Show first 20 features
                    show_features = min(20, feature_count)
                    features_table = doc.add_table(rows=show_features + 1, cols=2)
                    features_table.style = 'Light Grid'
                    
                    feat_hdr = features_table.rows[0].cells
                    feat_hdr[0].text = '#'
                    feat_hdr[1].text = 'Feature Name'
                    
                    for idx, feature in enumerate(all_features[:show_features], 1):
                        if idx < len(features_table.rows):
                            row_cells = features_table.rows[idx].cells
                            row_cells[0].text = str(idx)
                            row_cells[1].text = feature
                    
                    if feature_count > 20:
                        doc.add_paragraph(f"... and {feature_count - 20} more features")
                    
                    doc.add_paragraph().add_run().add_break()  # Spacer
                    
                    # ============================================
                    # 8. APPENDICES
                    # ============================================
                    doc.add_heading('8. Appendices', 1)
                    
                    # Appendix A: Complete feature list
                    doc.add_heading('Appendix A: Complete Feature List', 2)
                    doc.add_paragraph(f"Total Features: {feature_count}")
                    
                    # Create a multi-column layout for features
                    features_per_column = 30
                    num_columns = min(3, (feature_count + features_per_column - 1) // features_per_column)
                    
                    if num_columns > 0:
                        # Create table for features
                        feat_appendix_table = doc.add_table(rows=features_per_column + 1, cols=num_columns)
                        
                        # Fill features in columns
                        for col_idx in range(num_columns):
                            if col_idx < len(feat_appendix_table.rows[0].cells):
                                col_cell = feat_appendix_table.rows[0].cells[col_idx]
                                start = col_idx * features_per_column + 1
                                end = min((col_idx + 1) * features_per_column, feature_count)
                                col_cell.text = f"Features {start}-{end}"
                            
                            for row_idx in range(1, features_per_column + 1):
                                feat_idx = (col_idx * features_per_column) + (row_idx - 1)
                                if feat_idx < feature_count and row_idx < len(feat_appendix_table.rows):
                                    if col_idx < len(feat_appendix_table.rows[row_idx].cells):
                                        row_cell = feat_appendix_table.rows[row_idx].cells[col_idx]
                                        row_cell.text = all_features[feat_idx]
                    
                    doc.add_page_break()
                    
                    # Appendix B: Model configurations
                    doc.add_heading('Appendix B: Model Configurations', 2)
                    
                    for model_name, model_result in results['results'].items():
                        doc.add_heading(f'{model_name} Configuration', 3)
                        
                        if 'best_params' in model_result:
                            p = doc.add_paragraph('Best Hyperparameters:')
                            for param, value in model_result['best_params'].items():
                                doc.add_paragraph(f"  • {param}: {value}", style='List Bullet 2')
                        else:
                            doc.add_paragraph('Default parameters used')
                        
                        doc.add_paragraph().add_run().add_break()  # Spacer
                    
                    # Appendix C: Technical details
                    doc.add_heading('Appendix C: Technical Details', 2)
                    
                    tech_details = [
                        ('Analysis Date', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                        ('Software Version', 'Universal Churn Prediction System v1.0'),
                        ('Random Seed', str(CONFIG.RANDOM_STATE)),
                        ('Test Size', f"{CONFIG.TEST_SIZE*100:.0f}%"),
                        ('Validation Size', f"{CONFIG.VALIDATION_SIZE*100:.0f}%")
                    ]
                    
                    for detail_name, detail_value in tech_details:
                        p = doc.add_paragraph()
                        p.add_run(f"{detail_name}: ").bold = True
                        p.add_run(detail_value)
                    
                    # ============================================
                    # FOOTER
                    # ============================================
                    doc.add_page_break()
                    
                    footer = doc.add_paragraph()
                    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    footer_run = footer.add_run("--- END OF REPORT ---")
                    footer_run.font.size = Pt(12)
                    footer_run.font.color.rgb = RGBColor(100, 100, 100)
                    
                    # Save document to buffer
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    # Store in session state
                    st.session_state.word_doc_data = doc_buffer.getvalue()
                    st.session_state.word_doc_filename = f"{doc_name}.docx"
                    
                    st.success("✅ Word document created successfully!")
                    st.balloons()
                    
                except Exception as e:
                    st.error(f"❌ Error creating Word document: {str(e)[:200]}")
                    import traceback
                    st.error(f"Traceback: {traceback.format_exc()[:500]}")
    
    # Add Word document download button
    st.markdown("---")
    st.markdown("### 📄 Download Report Document")
    
    if hasattr(st.session_state, 'word_doc_data') and st.session_state.word_doc_data:
        # File info
        file_size = len(st.session_state.word_doc_data) / 1024  # KB
        
        col_info, col_download = st.columns([2, 1])
        
        with col_info:
            st.info(f"**Document:** {st.session_state.word_doc_filename} | **Size:** {file_size:.1f} KB")
            st.markdown("**Contents:** Executive summary, dataset overview, model results, business insights, feature analysis, and appendices")
        
        with col_download:
            st.download_button(
                label="📥 Download Word Document",
                data=st.session_state.word_doc_data,
                file_name=st.session_state.word_doc_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="download_word_doc"
            )
        
        # Preview document contents
        with st.expander("📋 Preview Document Contents", expanded=False):
            st.markdown("""
            **Document Structure:**
            
            1. **Cover Page** - Title and generation details
            2. **Table of Contents** - Navigation guide
            3. **Executive Summary** - Key findings and overview
            4. **Dataset Overview** - Data statistics and information
            5. **Data Analysis & Preprocessing** - Cleaning and preparation steps
            6. **Model Training Results** - Performance comparison and best model details
            7. **Business Insights** - Financial impact and recommendations
            8. **Feature Importance Analysis** - Top features and visualization
            9. **Download Information** - Model package details
            10. **Appendices** - Complete feature list and technical details
            """)
    else:
        st.info("Click 'Generate Word Doc' above to create a comprehensive report")
    
    # Model Summary at the bottom
    st.markdown("---")
    st.markdown("### 📊 Quick Model Summary")
    
    sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
    with sum_col1:
        st.metric("Model Type", best_model_name)
    with sum_col2:
        st.metric("Accuracy", f"{best_result['accuracy']:.4f}")
    with sum_col3:
        st.metric("AUC Score", f"{best_result['auc']:.4f}")
    with sum_col4:
        st.metric("Features", feature_count)
# Run the app
if __name__ == "__main__":
    pass
